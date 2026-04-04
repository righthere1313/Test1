import json
import sys
import time
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from fastapi.testclient import TestClient  # noqa: E402

from app.main import app  # noqa: E402


def main() -> int:
    run_id = str(int(time.time()))
    client = TestClient(app)

    health = client.get("/health")
    if health.status_code != 200:
        raise SystemExit(f"health failed: {health.status_code} {health.text}")

    resp = client.post(
        "/api/v1/generate/docx/auto",
        json={
            "topic": "勾股定理",
            "title": f"勾股定理教案-{run_id}",
        },
    )
    if resp.status_code != 200:
        raise SystemExit(f"docx auto failed: {resp.status_code} {resp.text}")
    data = resp.json()
    url = data.get("download_url")
    if not isinstance(url, str) or not url.startswith("/"):
        raise SystemExit(f"download_url invalid: {data}")

    dl = client.get(url)
    if dl.status_code != 200:
        raise SystemExit(f"download failed: {dl.status_code} {dl.text}")

    out = Path(__file__).resolve().parent / f"docx_auto_smoke_{run_id}.docx"
    out.write_bytes(dl.content)

    doc = Document(str(out))
    text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
    required = ["教学目标", "教学过程", "教学方法", "课堂活动设计", "课后作业"]
    missing = [x for x in required if x not in text]
    if missing:
        raise SystemExit(json.dumps({"ok": False, "missing": missing, "out": str(out)}, ensure_ascii=False))
    print(json.dumps({"ok": True, "out": str(out), "paragraphs": len(doc.paragraphs)}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

