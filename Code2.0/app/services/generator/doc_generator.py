from docx import Document
from app.schemas.generation import DocxDocument, DocxElement
import os
import uuid

class DocxGenerator:
    OUTPUT_DIR = "data/generated/docx"

    def __init__(self):
        os.makedirs(self.OUTPUT_DIR, exist_ok=True)

    def generate(self, data: DocxDocument) -> str:
        """
        Generates a DOCX file from the provided data.
        Returns the path to the generated file.
        """
        doc = Document()

        # 1. Add Title
        doc.add_heading(data.title, 0)

        # 2. Add Elements
        for element in data.elements:
            if element.type == "heading":
                doc.add_heading(element.content, level=element.level)
            elif element.type == "paragraph":
                doc.add_paragraph(element.content)
            elif element.type == "bullet":
                doc.add_paragraph(element.content, style='List Bullet')
            elif element.type == "numbered":
                doc.add_paragraph(element.content, style='List Number')
            else:
                # Default to paragraph
                doc.add_paragraph(element.content)

        # 3. Save File
        filename = f"{uuid.uuid4()}.docx"
        file_path = os.path.join(self.OUTPUT_DIR, filename)
        doc.save(file_path)

        return file_path
